# -*- coding: utf-8 -*-
"""
把 工程报告.md 转成格式美观的 Word 文档 工程报告.docx
用法：python code/deliverables/make_docx.py
依赖：python-docx（已装 1.2.0）
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.normpath(os.path.join(BASE_DIR, "..", "..", "工程报告.md"))
OUT = os.path.join(BASE_DIR, "工程报告.docx")

# ---------------- 字体 / 样式常量 ----------------
BODY_FONT = "宋体"            # 正文中文
HEAD_FONT = "微软雅黑"         # 标题中文
CODE_FONT = "Consolas"        # 代码 / 路径等宽
LATIN_FONT = "Times New Roman"  # 正文西文
NAVY = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)
HEADER_FILL = "D9E2F3"        # 表头底色（浅蓝）
CODE_FILL = "F2F2F2"          # 代码块底色（浅灰）

# 行内标记：**加粗** 或 `等宽`
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


# ---------------- 基础运行格式化 ----------------
def set_run(run, latin=LATIN_FONT, east=BODY_FONT, size=12, bold=None,
            italic=None, color=None):
    """统一设置 run 字体：西文 + 中文（eastAsia）+ 字号/粗斜体/颜色"""
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_run(paragraph, text, size=12, latin=LATIN_FONT, east=BODY_FONT,
            bold=False, color=None):
    run = paragraph.add_run(text)
    set_run(run, latin=latin, east=east, size=size, bold=bold, color=color)
    return run


def add_rich(paragraph, text, size=12, force_bold=False):
    """把一段文本按 **bold** / `code` 切成多个 run 写入段落"""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            add_run(paragraph, text[pos:m.start()], size=size,
                    bold=force_bold or False)
        tok = m.group(0)
        if tok.startswith("**"):
            add_run(paragraph, tok[2:-2], size=size, bold=True)
        else:
            add_run(paragraph, tok[1:-1], size=size, latin=CODE_FONT,
                    east=BODY_FONT, bold=force_bold)
        pos = m.end()
    if pos < len(text):
        add_run(paragraph, text[pos:], size=size, bold=force_bold)


# ---------------- 文档骨架 ----------------
def setup_document(doc):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    # Normal 默认样式：正文小四 12pt，中文宋体
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=22, latin=HEAD_FONT, east=HEAD_FONT, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    add_run(p, text, size=11, latin=HEAD_FONT, east=HEAD_FONT,
            color=GRAY)
    return p


def add_heading(doc, text, level):
    """level: 1 -> Heading 1, 2 -> Heading 2, 3 -> Heading 3"""
    h = doc.add_heading(level=level)
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    add_run(h, text, size=sizes.get(level, 12), latin=HEAD_FONT,
            east=HEAD_FONT, bold=True)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    if level >= 3:
        h.paragraph_format.keep_with_next = True
    return h


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    add_rich(p, text, size=12)
    return p


def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(0.4)
        # 浅灰底纹
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_FILL)
        pPr.append(shd)
        run = add_run(p, line if line else " ", size=10.5, latin=CODE_FONT,
                      east=BODY_FONT)
        # 保留行首空格
        t = run._r.find(qn("w:t"))
        if t is not None:
            t.set(qn("xml:space"), "preserve")


def add_bullet_list(doc, items):
    for depth, text in items:
        p = doc.add_paragraph(style="List Bullet")
        if depth > 0:
            p.paragraph_format.left_indent = Cm(1.0 + 0.6 * depth)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        add_rich(p, text, size=12)


# ---------------- 有序列表编号（每个列表从 1 重新开始） ----------------
_num_state = {"nid": None, "aid": None}


def _numbering_el(doc):
    return doc.part.numbering_part.element


def _next_ids(doc):
    if _num_state["nid"] is None:
        el = _numbering_el(doc)
        max_nid, max_aid = 0, 0
        for num in el.findall(qn("w:num")):
            max_nid = max(max_nid, int(num.get(qn("w:numId"))))
        for absn in el.findall(qn("w:abstractNum")):
            max_aid = max(max_aid, int(absn.get(qn("w:abstractNumId"))))
        _num_state["nid"], _num_state["aid"] = max_nid, max_aid
    _num_state["nid"] += 1
    _num_state["aid"] += 1
    return _num_state["nid"], _num_state["aid"]


def _build_abstract_num(abs_id):
    abs_el = OxmlElement("w:abstractNum")
    abs_el.set(qn("w:abstractNumId"), str(abs_id))
    for lvl in range(9):
        lvl_el = OxmlElement("w:lvl")
        lvl_el.set(qn("w:ilvl"), str(lvl))
        lvl_el.set(qn("w:tplc"), "0409000F")
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
        lvl_el.append(start)
        numFmt = OxmlElement("w:numFmt"); numFmt.set(qn("w:val"), "decimal")
        lvl_el.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "%" + str(lvl + 1) + ".")
        lvl_el.append(lvlText)
        lvlJc = OxmlElement("w:lvlJc"); lvlJc.set(qn("w:val"), "left")
        lvl_el.append(lvlJc)
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(420 + lvl * 420))
        ind.set(qn("w:hanging"), "420")
        pPr.append(ind)
        lvl_el.append(pPr)
        abs_el.append(lvl_el)
    return abs_el


def _build_num(num_id, abs_id):
    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(num_id))
    absId = OxmlElement("w:abstractNumId")
    absId.set(qn("w:val"), str(abs_id))
    num_el.append(absId)
    lvlOverride = OxmlElement("w:lvlOverride")
    lvlOverride.set(qn("w:ilvl"), "0")
    startOverride = OxmlElement("w:startOverride")
    startOverride.set(qn("w:val"), "1")
    lvlOverride.append(startOverride)
    num_el.append(lvlOverride)
    return num_el


def add_ordered_list(doc, items):
    num_id, abs_id = _next_ids(doc)
    el = _numbering_el(doc)
    el.append(_build_abstract_num(abs_id))
    el.append(_build_num(num_id, abs_id))
    for depth, text in items:
        p = doc.add_paragraph(style="List Number")
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(depth))
        numPr.append(ilvl)
        numId = OxmlElement("w:numId"); numId.set(qn("w:val"), str(num_id))
        numPr.append(numId)
        pPr.append(numPr)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        add_rich(p, text, size=12)


# ---------------- 表格 ----------------
def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    s = s.replace("\\|", "\x00")  # 转义竖线占位
    cells = [c.replace("\x00", "|").strip() for c in s.split("|")]
    return cells


def is_separator(cells):
    return bool(cells) and all(re.match(r"^:?-{2,}:?$", c) for c in cells)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def add_table(doc, rows):
    data = [split_row(r) for r in rows]
    header = data[0]
    body = [r for r in data[1:] if not is_separator(r)]
    ncols = max(len(r) for r in data)

    table = doc.add_table(rows=len(body) + 1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头：加粗 + 底纹 + 垂直居中
    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich(p, cell_text, size=10.5, force_bold=True)
        shade_cell(cell, HEADER_FILL)
        cell_vertical_center(cell)

    # 数据行
    for i, row in enumerate(body):
        for j, cell_text in enumerate(row):
            cell = table.cell(i + 1, j)
            p = cell.paragraphs[0]
            add_rich(p, cell_text, size=10.5)

    # 表格后留空距
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    sp.paragraph_format.line_spacing = 1.0
    return table


# ---------------- 主解析 ----------------
def parse_md(doc, lines):
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        # 文档大标题
        if line.startswith("# "):
            add_title(doc, line[2:].strip())
            i += 1
            continue

        # 引用（副标题）
        if line.startswith("> "):
            add_subtitle(doc, line[2:].strip())
            i += 1
            continue

        # 标题（## -> Heading 1, ### -> Heading 2, #### -> Heading 3）
        m = re.match(r"^(#{2,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1)) - 1
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # 表格块
        if stripped.startswith("|") and stripped.endswith("|"):
            rows = []
            while (i < n and lines[i].strip().startswith("|")
                   and lines[i].strip().endswith("|")):
                rows.append(lines[i].strip())
                i += 1
            add_table(doc, rows)
            continue

        # 围栏代码块
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束围栏
            add_code_block(doc, code_lines)
            continue

        # 无序列表
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            items = []
            while i < n:
                l = lines[i]
                if not l.strip():
                    break
                m2 = re.match(r"^(\s*)[-*]\s+(.*)$", l)
                if not m2:
                    break
                items.append((len(m2.group(1)) // 2, m2.group(2).strip()))
                i += 1
            add_bullet_list(doc, items)
            continue

        # 有序列表
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            items = []
            while i < n:
                l = lines[i]
                if not l.strip():
                    break
                m2 = re.match(r"^(\s*)(\d+)\.\s+(.*)$", l)
                if not m2:
                    break
                items.append((len(m2.group(1)) // 2, m2.group(3).strip()))
                i += 1
            add_ordered_list(doc, items)
            continue

        # 普通段落
        add_body(doc, stripped)
        i += 1


# ---------------- 验证 ----------------
def verify(out):
    from docx import Document as D2
    d = D2(out)
    paras = d.paragraphs
    print("== 验证结果 ==")
    print("文件大小: {:,} 字节".format(os.path.getsize(out)))
    print("段落数: {:,}".format(len(paras)))
    print("表格数: {}".format(len(d.tables)))
    for idx, t in enumerate(d.tables, 1):
        print("  表格 {}: {} 行 x {} 列".format(idx, len(t.rows), len(t.columns)))
    h1 = [p.text for p in paras if p.style.name == "Heading 1"]
    h2 = [p.text for p in paras if p.style.name == "Heading 2"]
    h3 = [p.text for p in paras if p.style.name == "Heading 3"]
    print("Heading 1 数量: {} -> {}".format(len(h1), h1))
    print("Heading 2 数量: {} -> {}".format(len(h2), h2))
    print("Heading 3 数量: {}".format(len(h3)))
    print("标题段落:", paras[0].text)
    print("日期副标题:", paras[1].text)


def main():
    if not os.path.exists(SRC):
        raise SystemExit("找不到报告: " + SRC)
    with open(SRC, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()
    setup_document(doc)
    parse_md(doc, md_text.splitlines())
    doc.save(OUT)
    print("已生成:", OUT)
    verify(OUT)


if __name__ == "__main__":
    main()
